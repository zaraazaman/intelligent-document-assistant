import streamlit as st
import streamlit as st
import os
import json
import requests
from datetime import datetime
from langchain_ollama import OllamaLLM, OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
import PyPDF2
import docx
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from io import StringIO




st.set_page_config(
    page_title="Intelligent Document Assistant",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
body, .css-1d391kg, .stApp {
    background-color: #f9f9f9;
    color: #0f766e;
    font-family: 'Inter', sans-serif;
}
/* Hero section */
.hero-title {
    font-size: 4rem;
    font-weight: 900;
    text-align: center;
    color: #004d4d;
    margin-bottom: 0.5rem;
}
.hero-subtitle {
    font-size: 1.7 rem;
    text-align: center;
    max-width: 900px;
    margin: 0 auto ;
    line-height: 1.6;
    color: #666666;
}
.feature-card {
    background-color: #e0f7f4;
    padding: 1rem;
    border-radius: 10px;
    text-align: center;
    margin: 0.3rem;
    font-size: 0.9rem;
    color: #0f766e;
}

.chat-message {
    padding: 0.8rem;
    border-radius: 12px;
    margin-bottom: 0.8rem;
    max-width: 80%;
}
.chat-user {
    background-color: #e0f7f4;
    margin-left: auto;
    text-align: right;
}
.chat-ai {
    background-color: #d1f0ec;
    margin-right: auto;
    text-align: left;
}
.chat-question {
    font-weight: 600;
    margin-bottom: 0.3rem;
}
.chat-answer {
    line-height: 1.5;
}

.stButton>button {
    background-color: #0f766e;
    color: white;
    border-radius: 8px;
    padding: 0.4rem 0.8rem;
    font-weight: 600;
}        
.stButton>button_sample {
    background-color: #D1F0EC;
    color: white;
    border-radius: 8px;
    padding: 0.4rem 0.8rem;
    font-weight: 600;
}
.stButton>button:hover {
    background-color: #0d5c54;
}

@keyframes bounce {
  0%, 20%, 50%, 80%, 100% {transform: translateY(0);}
  40% {transform: translateY(-8px);}
  60% {transform: translateY(-4px);}
}
.hero-title {
    animation: bounce 3s ;
}


@keyframes fadeIn {
    0% {opacity: 0; transform: translateY(10px);}
    100% {opacity: 1; transform: translateY(0);}
}
.hero-subtitle {
    animation: fadeIn 1.5s ease-out forwards;
    animation-delay: 0.2s;
    opacity: 0; /* start invisible for fade-in */
}


.main-bullets li {
    opacity: 0;
    transform: translateY(10px);
    animation: fadeIn 0.2s ease-out forwards;
}
.main-bullets li:nth-child(1) { animation-delay: 0.6s; }
.main-bullets li:nth-child(2) { animation-delay: 0.8s; }
.main-bullets li:nth-child(3) { animation-delay: 1.0s; }
.main-bullets li:nth-child(4) { animation-delay: 1.2s; }
.main-bullets li:nth-child(5) { animation-delay: 1.4s; }    


@keyframes slideIn {
    0% {opacity: 0; transform: translateY(20px);}
    100% {opacity: 1; transform: translateY(0);}
}
.feature-card {
    opacity: 0; /* start invisible */
    animation: slideIn 0.6s ease-out forwards;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
    cursor: pointer; /* looks interactive */
}

.feature-card:nth-child(1) { animation-delay: 0.3s; }
.feature-card:nth-child(2) { animation-delay: 0.5s; }
.feature-card:nth-child(3) { animation-delay: 0.7s; }


[data-testid="stSidebar"] {
    background-color: #e0f7f4;  /* soft teal */
    padding: 1.5rem 1rem;
    border-radius: 12px;
    transition: all 0.3s ease;
}


[data-testid="stSidebar"] h2 {
    font-size: 1.6rem;        /* smaller, sleeker */
    font-weight: 700;
    color: #004d4d;
    text-align: left;
    margin-bottom: 1rem;
    letter-spacing: 0.4px;
    border-bottom: 2px solid #0f766e;
    padding-bottom: 0.6rem;
    animation: fadeIn 0.3s ease-out;
}

.url-card {
    background: #d1f0ec;
    padding: 0.7rem 1rem;
    border-radius: 12px;
    margin-bottom: 1rem;
    font-size: 0.95rem;
    color: #004d4d;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
}
.url-card:hover {
    transform: translateY(-3px) scale(1.01);
    box-shadow: 0 8px 15px rgba(0,0,0,0.12);
}

.upload-card ::placeholder,
.url-card ::placeholder {
    color: #006666;
    font-weight: 500;
    font-size: 0.9rem;
}

.file-info {
    background-color: #f0fdfc;
    padding: 0.6rem;
    border-radius: 8px;
    margin-bottom: 0.4rem;
    color: #0f766e;
    transition: transform 0.2s ease, box-shadow 0.2s ease;
}


.file-info:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 10px rgba(0,0,0,0.1);
}

[data-testid="stSidebar"] .stSelectbox {
    border-radius: 10px;
    transition: all 0.3s ease;
}
[data-testid="stSidebar"] .stSelectbox:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 12px rgba(0,0,0,0.15);
}        

.feature-card {
    transition: transform 0.3s ease, box-shadow 0.3s ease;
}
.feature-card:hover {
    transform: translateY(-5px);
    box-shadow: 0 8px 15px rgba(0,0,0,0.1);
}

.stButton>button:hover {
    transform: translateY(-2px);
    opacity: 0.9;
}

.file-info {
    background-color: #e0f7f4;
    padding: 0.6rem;
    border-radius: 6px;
    margin-bottom: 0.4rem;
    color: #0f766e;
}      
.sample-question {
    background-color: #D1F0EC;
    border: 1px solid #0f766e;
    padding: 0.6rem 1rem;
    border-radius: 10px;
    text-align: center;
    margin: 0.3rem;
    font-size: 0.9rem;
    color: #0f766e;
    cursor: pointer;
    display: inline-block;
    transition: all 0.2s ease-in-out;
}
.sample-question:hover {
    background-color: #b8e8e2;
    transform: translateY(-2px);
}
.footer {
    position: fixed;
    bottom: 0;
    left: 0;
    width: 100%;
    background-color: #f9f9f9; 
    text-align: center;
    padding: 10px;
    font-size: 0.85rem;
    color: #666;
    border-top: 1px solid #ddd;
}
</style>
""", unsafe_allow_html=True)

class DocAssistant:
    def __init__(self):
        self.llm = None
        self.embeddings = None
        self.vectorstore = None
        self.qa_chain = None
        self.documents = []
        self.conversation_history = []
        self.document_sources = []

    

    def initialize_models(self):
        if self.llm is None:
            try:
                self.llm = OllamaLLM(model="mistral:7b")
                self.embeddings = OllamaEmbeddings(model="mistral:7b")
                return True
            except Exception as e:
                st.error(f"Error initializing models: {str(e)}")
                return False
        return True

    
    
    def extract_text_from_pdf(self, pdf_file):
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text += f"[Page {page_num + 1}]\n{page_text}\n\n"
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return None


    
    def extract_text_from_docx(self, docx_file):
        try:
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading Word document: {str(e)}")
            return None


    
    def extract_text_from_url(self, url):
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            return text
        except Exception as e:
            st.error(f"Error extracting text from URL {url}: {str(e)}")
            return None

    
    
    def validate_url(self, url):
        try:
            result = urlparse(url)
            return all([result.scheme, result.netloc])
        except:
            return False


    
    def process_documents(self, uploaded_files=None, urls=None):
        if not self.initialize_models():
            return False
        documents = []
        self.document_sources = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        total_items = (len(uploaded_files) if uploaded_files else 0) + (len(urls) if urls else 0)
        processed_items = 0
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    text = None
                    file_type = uploaded_file.type
                    if file_type == "application/pdf":
                        text = self.extract_text_from_pdf(uploaded_file)
                    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        text = self.extract_text_from_docx(uploaded_file)
                    elif file_type == "text/plain":
                        text = str(uploaded_file.read(), "utf-8")
                    else:
                        st.warning(f"Unsupported file type: {file_type}")
                        continue
                    if text and text.strip():
                        doc = Document(
                            page_content=text,
                            metadata={
                                "source": uploaded_file.name,
                                "type": "file",
                                "size": uploaded_file.size
                            }
                        )
                        documents.append(doc)
                        self.document_sources.append({
                            "name": uploaded_file.name,
                            "type": "File",
                            "size": f"{uploaded_file.size / 1024:.1f} KB"
                        })
                except Exception as e:
                    st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                processed_items += 1
                progress_bar.progress(processed_items / total_items)
        
        if urls:
            for url in urls:
                url = url.strip()
                if url and self.validate_url(url):
                    try:
                        text = self.extract_text_from_url(url)
                        if text and text.strip():
                            doc = Document(
                                page_content=text,
                                metadata={
                                    "source": url,
                                    "type": "url"
                                }
                            )
                            documents.append(doc)
                            self.document_sources.append({
                                "name": url,
                                "type": "URL",
                                "size": f"{len(text) / 1024:.1f} KB"
                            })
                    except Exception as e:
                        st.error(f"Error processing URL {url}: {str(e)}")
                
                processed_items += 1
                progress_bar.progress(processed_items / total_items)
        if documents:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", " ", ""])
            splits = text_splitter.split_documents(documents)
            self.vectorstore = FAISS.from_documents(splits, self.embeddings)    
            
            prompt_template = """
            You are an intelligent document assistant. Use the provided context to answer the question accurately and concisely.
            Guidelines:
            - Provide clear, accurate answers based on the context
            - If information is not in the documents, say so clearly
            - Format your response appropriately for the question type
            - Include relevant details but keep responses focused
            Context:
            {context}
            Question:
            {input}
            Answer:
            """

            
            PROMPT = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "input"]
            )
            document_chain = create_stuff_documents_chain(
                llm=self.llm,
                prompt=PROMPT
            )
            self.qa_chain = create_retrieval_chain(
                retriever=self.vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}
                ),
                combine_docs_chain=document_chain
            )
            
            self.documents = documents
            progress_bar.progress(1.0)
            status_text.text("✔ Processed successfully!")
            return True
        
        status_text.text("⚠️ No documents were processed.")
        return False


    
    def ask_question(self, question):
        if not self.qa_chain:
            return "Please upload and process documents first."
        
        try:
            with st.spinner("Analyzing documents..."):
                response = self.qa_chain.invoke({"input": question})
                answer = response.get("answer", "No answer found")
                
                self.conversation_history.append({
                    "question": question,
                    "answer": answer,
                    "timestamp": datetime.now(),
                    "sources": [doc.metadata.get("source", "Unknown") for doc in response.get("context", [])]
                })
                
                return answer
        except Exception as e:
            return f"Error: {str(e)}"




def main():
    if 'assistant' not in st.session_state:
        st.session_state.assistant = DocAssistant()
    if 'documents_processed' not in st.session_state:
        st.session_state.documents_processed = False

    st.markdown('<div class="hero-title">📚 Intelligent Document Assistant</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("""<div class="hero-subtitle">Turn your documents into a smart knowledge base and get instant, AI-powered answers.</div>
                <br>
                <br>
                <div class="main-bullets">
                <ul>
                <li> Import PDFs, Word docs, text files, or even URLs in seconds</li>
                <li> Build a personalized, searchable database of all your docs</li>
                <li> Ask natural questions and get precise, context-aware answers</li>
                <li> Keep a history of uploads and revisit them whenever you need</li>
                <li> Download AI responses as text files with a single click</li>
                </ul>
                </div>
                <br>
                """, unsafe_allow_html=True)
    
    st.markdown("### Key Features")
    col1, col2, col3 = st.columns(3)
    features = [
        ("Multi-format Support", "PDFs, DOCX, TXT, URLs"),
        ("Smart Search", "Create searchable document indexes"),
        ("AI-Powered Q&A", "Context-aware answers from your content")
    ]
    for col, (title, desc) in zip([col1, col2, col3], features):
        col.markdown(f'<div class="feature-card"><strong>{title}</strong><br>{desc}</div>', unsafe_allow_html=True)
    
   
    if st.session_state.documents_processed:
        st.markdown("<br>", unsafe_allow_html=True)  
        st.subheader("💬 What would you like to know?")
        st.markdown("<br>", unsafe_allow_html=True)  
        st.markdown("**Quick Suggestions**",unsafe_allow_html=True)
        sample_questions = [
            "What are the main insights?",
            "Summarize the key points.",
            "What conclusions can be drawn?",
            "Explain the methodology clearly.",
            "List down important features."
            ]

        cols = st.columns(len(sample_questions))
        for i, q in enumerate(sample_questions):
            cols[i].button(
                q,
                key=f"sample_{i}",
                on_click=lambda q=q: st.session_state.update(question_input=q)
            )
        st.markdown("<br>", unsafe_allow_html=True) 
        question = st.text_input(
            "Your question:",
            placeholder="e.g., What are the key factors described",
            key="question_input"
        )

        col1, col2 = st.columns([3,3])  

        with col1:
            ask_button = st.button("Answer", type="primary", use_container_width=True)

        with col2:
            clear_button = st.button(
                "Clear",
                type="secondary",
                use_container_width=True,
                on_click=lambda: st.session_state.update(question_input="")
            )

        if ask_button and question:
            answer = st.session_state.assistant.ask_question(question)
            st.markdown(f"""
            <div class="chat-message chat-ai">
                <div class="chat-answer">{answer}</div>
                <div class="chat-timestamp">⏱︎  {datetime.now().strftime('%H:%M:%S')}</div>
            </div>
            """, unsafe_allow_html=True)

            export_chat=f"Question: {question}\n\nAnswer: {answer}"
            st.download_button(
                label="⎙ Export",
                data=export_chat,
                file_name="answer.txt",
                mime="text/plain"
            )


    with st.sidebar:
        st.markdown('<h2>Files & URLs</h2><br>', unsafe_allow_html=True)
        uploaded_files = st.file_uploader("🔗 Attach Documents", accept_multiple_files=True, type=['pdf','docx','txt'])
        url_input = st.text_area("🌐 Add URLs (one per line)", placeholder="https://doc.com/c/", height=80)
        urls = [url.strip() for url in url_input.split('\n') if url.strip()] if url_input else None
        if uploaded_files or urls:
            if st.button("Process",use_container_width=True):
                with st.spinner("Processing & creating search index"):
                    success = st.session_state.assistant.process_documents(uploaded_files, urls)
                    if success:
                        st.session_state.documents_processed = True
                        st.balloons()
                        st.rerun()

        st.markdown("### Uploaded Files")
        for doc in st.session_state.assistant.document_sources:
            st.markdown(f'<div class="file-info"><strong>{doc["name"]}</strong> | Type: {doc["type"]} | Size: {doc["size"]}</div>', unsafe_allow_html=True)

        if st.session_state.assistant.conversation_history:
            st.markdown("---")
            st.subheader("Chat History")
            options = [
                f"{i+1}. {conv['question'][:40]}..." 
                for i, conv in enumerate(reversed(st.session_state.assistant.conversation_history[-10:]))
            ]

            selected = st.selectbox(
                "Review past chats",
                options,
                index=None,
                placeholder=""
            )
            if selected:
                idx = int(selected.split(".")[0]) - 1
                convs = list(reversed(st.session_state.assistant.conversation_history[-10:]))
                st.session_state.selected_conv = convs[idx]
        else:
            st.markdown("")

    if "selected_conv" in st.session_state and st.session_state.selected_conv:
        conv = st.session_state.selected_conv
        st.markdown("### 📌 Selected Chat")
        st.markdown(f"""
        <div class="chat-message">
            <div class="chat-question"><b>Q:</b> {conv['question']}</div>
            <div class="chat-answer"><b>A:</b> {conv['answer']}</div>
            
        </div>
        """, unsafe_allow_html=True)    
        
    st.markdown("""
                <div class="footer">
<span>Intelligent Document Assistant</span> • 
        <span>Powered by <b>Streamlit</b> & <b>LangChain</b></span> •
        <span>© 2025 Intelligent Document Assistant. All rights reserved.</span></div>""",unsafe_allow_html=True)
        
if __name__ == "__main__":
    main()
